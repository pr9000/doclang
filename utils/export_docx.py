#!/usr/bin/env python3

"""
Export spec.md to Word (DOCX).

Features:
- Parses markdown headings, paragraphs, lists, code blocks, tables, images.
- Embeds images referenced in the markdown into the DOCX.
- Inserts a Word Table of Contents field (updates on open in Word).
- Optionally uses exports/templates/word.dotx if present.

Dependencies:
- python-docx (pip install python-docx)

Note: This is a lightweight Markdown parser tailored for this repository's
spec.md. It aims for faithful structure rather than exhaustive Markdown
feature coverage. If you need more fidelity, consider running Pandoc externally
and post-processing to add a TOC field.
"""

from __future__ import annotations

import argparse
import os
import shutil
import re
import sys
from pathlib import Path
from typing import List, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE
except ImportError as exc:
    print(
        "Missing dependency: python-docx. Install with `pip install python-docx`.",
        file=sys.stderr,
    )
    raise


ROOT = Path(__file__).resolve().parents[1]
EXPORTS_DIR = ROOT / "exports"
WORD_TEMPLATE = EXPORTS_DIR / "templates" / "word.dotx"
DEFAULT_MD = ROOT / "spec.md"
OUTPUT_DOCX = EXPORTS_DIR / "doclang.docx"


def add_toc(document: Document) -> None:
    # Insert a Word ToC field: TOC \o "1-3" \h \z \u
    p = document.add_paragraph()
    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instr)

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_separate)

    # Placeholder text shown before updating fields in Word.
    hint = p.add_run("Table of Contents – update to populate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    hint._r.append(fld_end)



def process_html_paragraph(document: Document, text: str) -> None:
    """Process paragraph text that contains HTML elements like <ul>, <li>, and <br>.

    Splits the text into segments and processes each appropriately:
    - Text before/after HTML elements as regular paragraphs
    - <ul>...</ul> as bulleted lists
    - <br> or <br/> as line breaks within paragraphs
    """
    # Pattern to match <ul>...</ul> blocks
    ul_pattern = re.compile(r'<ul>(.*?)</ul>', re.DOTALL | re.IGNORECASE)
    # Pattern to match <li>...</li> items
    li_pattern = re.compile(r'<li>(.*?)</li>', re.DOTALL | re.IGNORECASE)

    cursor = 0

    # Find all <ul> blocks
    for ul_match in ul_pattern.finditer(text):
        # Process text before the <ul>
        if ul_match.start() > cursor:
            pre_text = text[cursor:ul_match.start()].strip()
            if pre_text:
                # Process <br> tags in the pre-text
                process_br_tags(document, pre_text)

        # Process the <ul> content
        ul_content = ul_match.group(1)
        for li_match in li_pattern.finditer(ul_content):
            li_text = li_match.group(1).strip()
            # HTML unescape but preserve markdown formatting (backticks, etc.)
            li_text = _html_unescape(li_text)
            p = document.add_paragraph(style="List Bullet")
            # Process markdown formatting (inline code, links, bold, italic)
            _add_inline_formatted_runs(p, li_text)

        cursor = ul_match.end()

    # Process remaining text after last <ul>
    if cursor < len(text):
        post_text = text[cursor:].strip()
        if post_text:
            process_br_tags(document, post_text)


def process_br_tags(document: Document, text: str) -> None:
    """Process text containing <br> tags by splitting into multiple runs with line breaks."""
    br_pattern = re.compile(r'<br\s*/?>', re.IGNORECASE)

    # Split by <br> tags
    segments = br_pattern.split(text)

    if len(segments) == 1:
        # No <br> tags, process as normal paragraph
        add_formatted_paragraph(document, text)
    else:
        # Create a single paragraph with line breaks
        p = document.add_paragraph()
        for idx, segment in enumerate(segments):
            segment = segment.strip()
            if segment:
                _add_inline_formatted_runs(p, segment)
            # Add line break after each segment except the last
            if idx < len(segments) - 1:
                p.add_run().add_break(WD_BREAK.LINE)

def finalize_paragraph_buf(document: Document, buf: List[str]) -> None:
    if not buf:
        return
    text = " ".join(line.strip() for line in buf).strip()
    if text:
        # Check if text contains HTML lists or line breaks that should be processed
        if '<ul>' in text or '<br' in text:
            process_html_paragraph(document, text)
        else:
            add_formatted_paragraph(document, text)
    buf.clear()


def replace_docx_parts_from_template(template_path: Path, docx_path: Path, parts: List[str]) -> None:
    """
    Copy specific parts from a DOCX/DOTX template into a target DOCX, replacing
    any existing entries with the same names. This avoids duplicate entries and
    the associated zipfile UserWarning messages.

    - template_path: path to a .docx/.dotx file to read parts from
    - docx_path: path to the .docx file to modify in-place
    - parts: list of ZIP member names to copy (e.g., 'word/styles.xml')
    """
    import zipfile

    parts_set = set(parts)
    tmp_path = docx_path.with_suffix(docx_path.suffix + ".tmp")

    # Build a new zip: copy everything from the destination except the parts
    # we intend to replace, then write the parts from the template.
    with zipfile.ZipFile(docx_path, "r") as zread, zipfile.ZipFile(tmp_path, "w") as zwrite:
        # Copy existing entries except those to be replaced
        for info in zread.infolist():
            if info.filename in parts_set:
                continue
            data = zread.read(info.filename)
            # Preserve metadata by reusing ZipInfo when writing
            zwrite.writestr(info, data)

        # Append parts from the template if present
        with zipfile.ZipFile(template_path, "r") as zsrc:
            for name in parts:
                try:
                    src_info = zsrc.getinfo(name)
                    data = zsrc.read(name)
                    zwrite.writestr(src_info, data)
                except KeyError:
                    # Part not present in template; skip
                    pass

    # Atomically replace the original docx with the rebuilt archive
    tmp_path.replace(docx_path)


_LIST_BULLET_RE = re.compile(r"^(?P<indent>\s*)([-+*])\s+(?P<text>.+?)\s*$")
_LIST_ORDERED_RE = re.compile(r"^(?P<indent>\s*)(\d+)([.)])\s+(?P<text>.+?)\s*$")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*?)\s*$")
_IMAGE_RE = re.compile(r'!\[(?P<alt>[^\]]*)\]\((?P<src>[^)\s]+)(?:\s+"(?P<title>[^"]*)")?\)')
_LINK_OR_URL_RE = re.compile(
    r'(?:\[(?P<link_text>[^\]]+)\]\((?P<link_url>[^)\s]+)(?:\s+"[^"]*")?\))|(?P<bare>(?:https?://|mailto:)[^\s<)]+)'
)


def _strip_markdown_formatting(text: str) -> str:
    """Strip markdown formatting from text (inline code, links, bold, italic).

    Converts:
    - `code` -> code
    - [text](url) -> text
    - **bold** -> bold
    - *italic* -> italic
    """
    # Remove inline code backticks
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # Convert markdown links to just the link text
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    # Remove bold markers, then italic (order matters)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*([^*\n]+?)\*(?!\*)", r"\1", text)
    return text


def parse_table_block(lines: List[str], start: int) -> Tuple[int, List[List[str]], Optional[List[str]]]:
    # Detect a GitHub-style pipe table starting at index start.
    # Returns (end_index_exclusive, rows, aligns)
    table_lines = []
    i = start
    while i < len(lines) and "|" in lines[i].strip():
        table_lines.append(lines[i].strip())
        i += 1
        # Stop if we hit a blank line.
        if i < len(lines) and not lines[i].strip():
            break
    if len(table_lines) < 2:
        return start, [], None
    # Second line must be the separator with dashes and optional colons.
    header = [c.strip() for c in split_md_row(table_lines[0])]
    sep = [c.strip() for c in split_md_row(table_lines[1])]
    if not all(re.match(r"^:?-{3,}:?$", s) for s in sep) or len(sep) != len(header):
        return start, [], None

    aligns: List[str] = []
    for s in sep:
        if s.startswith(":") and s.endswith(":"):
            aligns.append("center")
        elif s.endswith(":"):
            aligns.append("right")
        elif s.startswith(":"):
            aligns.append("left")
        else:
            aligns.append("left")

    rows = [header]
    for r in table_lines[2:]:
        cells = [c.strip() for c in split_md_row(r)]
        # Normalize cell count
        while len(cells) < len(header):
            cells.append("")
        rows.append(cells[: len(header)])
    return i, rows, aligns


def split_md_row(row: str) -> List[str]:
    # Splits a pipe table row into cells, ignoring outer pipes.
    s = row.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.replace("\\|", "|") for c in re.split(r"\s*\|\s*", s)]


def add_code_block(document: Document, code_lines: List[str], language: Optional[str]) -> None:
    # Add a monospaced, preformatted code block.
    p = document.add_paragraph()
    # Light grey background for the whole block
    try:
        _set_paragraph_shading(p, fill="F2F2F2")
    except Exception:
        pass
    run = p.add_run()
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    for idx, line in enumerate(code_lines):
        if idx:
            run.add_break(WD_BREAK.LINE)
        run.add_text(line.rstrip("\n"))


INLINE_CODE_TAG_RE = re.compile(r"<inline-code>(.*?)</inline-code>")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
ITALIC_RE = re.compile(r"(?<!\*)\*([^*\n]+?)\*(?!\*)")
CODE_SPAN_RE = re.compile(r"(<inline-code>(.*?)</inline-code>|`([^`]+)`)")


def _set_paragraph_shading(paragraph, fill: str = "F2F2F2") -> None:
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _set_run_shading(run, fill: str = "F2F2F2") -> None:
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    rPr.append(shd)


def _add_italic_runs(paragraph, text: str) -> None:
    """Add runs for *italic* segments (single asterisks, not **bold**)."""
    pos = 0
    for m in ITALIC_RE.finditer(text):
        if m.start() > pos:
            t = text[pos : m.start()]
            if t:
                paragraph.add_run(t)
        ir = paragraph.add_run(m.group(1))
        ir.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_bold_and_text_runs(paragraph, text: str) -> None:
    """Add runs for plain text, *italic*, and **bold** (bold processed before italic)."""
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            _add_italic_runs(paragraph, text[pos : m.start()])
        br = paragraph.add_run(m.group(1))
        br.bold = True
        pos = m.end()
    if pos < len(text):
        _add_italic_runs(paragraph, text[pos:])


def _add_links_and_text_runs(paragraph, text: str) -> None:
    """Add runs handling Markdown links [text](url) and bare URLs, plus bold/italic.

    This does not handle inline code; the caller should have already split
    code spans and pass only normal text here.

    Anchor links like [text](#anchor) become in-document hyperlinks to bookmarks.
    """
    pos = 0
    for m in _LINK_OR_URL_RE.finditer(text):
        if m.start() > pos:
            pre = text[pos : m.start()]
            if pre:
                _add_bold_and_text_runs(paragraph, pre)
        if m.group("link_url") is not None:
            link_text = m.group("link_text")
            link_url = m.group("link_url")
            if link_url.startswith("#"):
                add_internal_hyperlink(paragraph, link_url, link_text)
            else:
                add_hyperlink(paragraph, link_url, link_text)
        else:
            url = m.group("bare")
            add_hyperlink(paragraph, url, url)
        pos = m.end()
    if pos < len(text):
        _add_bold_and_text_runs(paragraph, text[pos:])


def _add_inline_formatted_runs(paragraph, text: str) -> None:
    """Process inline formatting: links (with embedded code), code spans, bold, italic.

    Links are processed first to handle cases like [`<element>`](#anchor) where
    the link text contains inline code.
    """
    # First pass: extract and process markdown links, preserving their structure
    cursor = 0
    # Segments: ("plain", text) | ("url", url, text) | ("anchor", anchor, text)
    segments: List[Tuple[str, ...]] = []

    for m in _LINK_OR_URL_RE.finditer(text):
        if m.start() > cursor:
            segments.append(("plain", text[cursor:m.start()]))

        if m.group("link_url") is not None:
            link_text = m.group("link_text")
            link_url = m.group("link_url")
            if link_url.startswith("#"):
                segments.append(("anchor", link_url, link_text))
            else:
                segments.append(("url", link_url, link_text))
        else:
            url = m.group("bare")
            segments.append(("url", url, url))

        cursor = m.end()

    if cursor < len(text):
        segments.append(("plain", text[cursor:]))

    for segment in segments:
        kind = segment[0]
        if kind == "plain":
            _process_code_and_bold(paragraph, segment[1])
        elif kind == "anchor":
            add_internal_hyperlink(paragraph, segment[1], segment[2])
        else:
            url, link_text = segment[1], segment[2]
            temp_p = paragraph._element.getparent().makeelement(paragraph._element.tag)
            temp_para = paragraph.__class__(temp_p, paragraph._parent)
            _process_code_and_bold(temp_para, link_text)
            if temp_para.runs:
                link_text = "".join(r.text or "" for r in temp_para.runs)
            add_hyperlink(paragraph, url, link_text)


def _process_code_and_bold(paragraph, text: str) -> None:
    """Process text for inline code spans, bold, and italic, adding runs to paragraph."""
    cursor = 0
    for m in CODE_SPAN_RE.finditer(text):
        if m.start() > cursor:
            normal_text = text[cursor : m.start()]
            if normal_text:
                _add_bold_and_text_runs(paragraph, normal_text)
        code_txt = m.group(2) if m.group(2) is not None else m.group(3)
        r = paragraph.add_run(code_txt)
        r.font.name = "Consolas"
        try:
            _set_run_shading(r, fill="F2F2F2")
        except Exception:
            pass
        cursor = m.end()
    if cursor < len(text):
        tail = text[cursor:]
        _add_bold_and_text_runs(paragraph, tail)


def add_formatted_paragraph(document: Document, text: str):
    p = document.add_paragraph()
    _add_inline_formatted_runs(p, text)
    return p


def add_image(document: Document, src: str, base_dir: Path) -> bool:
    # Resolve and add image; return True if added.
    # Skip if remote URL.
    if re.match(r"^[a-z]+://", src):
        return False
    img_path = (base_dir / src).resolve()
    if not img_path.exists():
        # Try relative to project root as fallback.
        alt_path = (ROOT / src).resolve()
        if alt_path.exists():
            img_path = alt_path
        else:
            print(f"[warn] Image not found: {src}", file=sys.stderr)
            return False
    try:
        # Resize to visible page width while preserving aspect ratio.
        try:
            section = document.sections[-1]
            max_width = section.page_width - section.left_margin - section.right_margin
            document.add_picture(str(img_path), width=max_width)
        except Exception:
            # Fallback if section metrics are unavailable
            document.add_picture(str(img_path))
        return True
    except Exception as e:
        print(f"[warn] Failed to embed image {src}: {e}", file=sys.stderr)
        return False


FIGURE_IMG_SRC_RE = re.compile(r"<img[^>]*\bsrc\s*=\s*['\"]([^'\"]+)['\"][^>]*>", re.IGNORECASE | re.DOTALL)
FIGCAPTION_RE = re.compile(r"<figcaption[^>]*>(.*?)</figcaption>", re.IGNORECASE | re.DOTALL)
HTML_A_TAG_RE = re.compile(r"<a[^>]*\bhref\s*=\s*['\"]([^'\"]+)['\"][^>]*>(.*?)</a>", re.IGNORECASE | re.DOTALL)
HTML_TAG_RE = re.compile(r"<[^>]+>")
from html import unescape as _html_unescape

def markdown_anchor(text: str) -> str:
    """Slug for in-document cross-references (aligned with generate_reference.py)."""
    return text.replace("`", "").replace("<", "").replace(">", "").lower().replace(" ", "-")


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    """Insert a Word bookmark on a paragraph for internal hyperlink targets."""
    p = paragraph._p
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    p.insert(0, start)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    p.append(end)


def _style_hyperlink_runs(hyperlink) -> None:
    for run in hyperlink.findall(qn("w:r")):
        r_pr = run.find(qn("w:rPr"))
        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            run.insert(0, r_pr)
        r_style = OxmlElement("w:rStyle")
        r_style.set(qn("w:val"), "Hyperlink")
        r_pr.append(r_style)


def add_internal_hyperlink(paragraph, anchor: str, text: str) -> None:
    """Add a clickable in-document hyperlink to a bookmark (e.g. [text](#anchor))."""
    anchor_name = anchor.lstrip("#")
    mark = len(paragraph._p)
    _process_code_and_bold(paragraph, text)
    moved = paragraph._p[mark:]
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor_name)
    for child in list(moved):
        paragraph._p.remove(child)
        hyperlink.append(child)
    _style_hyperlink_runs(hyperlink)
    paragraph._p.append(hyperlink)


def add_hyperlink(paragraph, url: str, text: str):
    """Add a clickable hyperlink run to a paragraph.

    Uses low-level OXML since python-docx lacks a high-level API.
    """
    # Create relationship id
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and set relationship id
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a run with the hyperlink style
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)

    w_t = OxmlElement("w:t")
    w_t.text = text
    new_run.append(w_t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph


def add_heading_with_bookmark(document: Document, raw_title: str, level: int, bookmark_id: int) -> None:
    """Add a heading paragraph and a bookmark for cross-reference targets."""
    display = _strip_markdown_formatting(raw_title)
    paragraph = document.add_heading(display, level=level)
    add_bookmark(paragraph, markdown_anchor(raw_title), bookmark_id)


def _strip_html_comments_outside_code(md_text: str) -> str:
    """Remove HTML comments (<!-- ... -->) outside fenced code blocks.

    Preserves comment markers that appear inside fenced code blocks. Also
    removes multi-line comments and inline fragments within a line.
    """
    lines = md_text.splitlines(keepends=True)
    out: List[str] = []
    in_code = False
    code_fence: Optional[str] = None
    in_comment = False
    for line in lines:
        # Detect fenced code block start/end (must be fence-only line, like parser)
        if not in_comment:
            mcode = re.match(r"^(\s*)(`{3,}|~{3,})(\w+)?\s*$", line)
            if mcode:
                # Toggle code state
                if not in_code:
                    in_code = True
                    code_fence = mcode.group(2)
                else:
                    # Closing fence (any fence-only line ends block in our parser)
                    in_code = False
                    code_fence = None
                out.append(line)
                continue

        if in_code:
            out.append(line)
            continue

        # Strip HTML comments in non-code text
        i = 0
        L = len(line)
        buf: List[str] = []
        while i < L:
            if in_comment:
                end = line.find("-->", i)
                if end == -1:
                    # Entire remainder is within comment
                    i = L
                    in_comment = True
                    break
                else:
                    # Close comment and continue scanning after it
                    i = end + 3
                    in_comment = False
                    continue
            else:
                start = line.find("<!--", i)
                if start == -1:
                    buf.append(line[i:])
                    break
                else:
                    # Keep text before the comment start
                    if start > i:
                        buf.append(line[i:start])
                    i = start + 4
                    in_comment = True
                    continue
        out.append("".join(buf))
    return "".join(out)


def parse_markdown_to_docx(document: Document, md_text: str, base_dir: Path) -> None:
    lines = md_text.splitlines()
    i = 0
    para_buf: List[str] = []
    # Track the last seen Markdown heading level to place sub-headers
    current_heading_level: int = 1
    next_bookmark_id = 0

    # For list indentation, 2 spaces per level is common in this repo.
    def list_level_for_indent(s: str) -> int:
        return max(0, len(s.replace("\t", "  ")) // 2)

    while i < len(lines):
        line = lines[i]

        # Handle <details> wrapper with <summary> header: inject a sub-heading and
        # skip wrapper/comment lines; keep inner content for normal parsing.
        if re.match(r"^\s*<details>\s*$", line):
            finalize_paragraph_buf(document, para_buf)
            i += 1
            # Optional immediate <summary>Title</summary>
            if i < len(lines):
                msum = re.match(r"^\s*<summary>(.*?)</summary>\s*$", lines[i])
                if msum:
                    summary_text = msum.group(1).strip()
                    sub_level = min(current_heading_level + 1, 6)
                    add_heading_with_bookmark(
                        document, summary_text, sub_level, next_bookmark_id
                    )
                    next_bookmark_id += 1
                    # Do not update current_heading_level so siblings remain at same level
                    i += 1
            # Skip optional blank/comment lines after <summary>
            while i < len(lines):
                s = lines[i].strip()
                if not s or re.match(r"^<!--.*?-->$", s):
                    i += 1
                    continue
                break
            continue

        # Ignore closing </details>
        if re.match(r"^\s*</details>\s*$", line):
            finalize_paragraph_buf(document, para_buf)
            i += 1
            continue

        # HTML <figure> block with one or more <img> tags: embed images like markdown
        if re.match(r"^\s*<figure\b", line, flags=re.IGNORECASE):
            finalize_paragraph_buf(document, para_buf)
            block_lines = [line]
            i += 1
            # Accumulate until closing </figure>
            while i < len(lines) and re.search(r"</figure>", lines[i], flags=re.IGNORECASE) is None:
                block_lines.append(lines[i])
                i += 1
            if i < len(lines):
                block_lines.append(lines[i])
                i += 1
            block = "\n".join(block_lines)
            # Find and embed all images within the figure
            for src in FIGURE_IMG_SRC_RE.findall(block):
                add_image(document, src.strip(), base_dir)
            # If a figcaption exists, add it as a caption paragraph beneath the image(s)
            mcap = FIGCAPTION_RE.search(block)
            if mcap:
                caption_html = mcap.group(1).strip()
                # Build a caption paragraph. Try to use the 'Caption' style if available.
                try:
                    p = document.add_paragraph(style="Caption")
                except Exception:
                    p = document.add_paragraph()
                # Mix plain text and HTML <a href> links
                pos = 0
                for am in HTML_A_TAG_RE.finditer(caption_html):
                    if am.start() > pos:
                        pre_html = caption_html[pos : am.start()]
                        pre_txt = _html_unescape(HTML_TAG_RE.sub("", pre_html))
                        if pre_txt:
                            p.add_run(pre_txt)
                    url = _html_unescape(am.group(1))
                    link_text_html = am.group(2)
                    link_text = _html_unescape(HTML_TAG_RE.sub("", link_text_html)) or url
                    add_hyperlink(p, url, link_text)
                    pos = am.end()
                if pos < len(caption_html):
                    tail_html = caption_html[pos:]
                    tail_txt = _html_unescape(HTML_TAG_RE.sub("", tail_html))
                    if tail_txt:
                        p.add_run(tail_txt)
                # Light italic for readability if no Caption style applied
                try:
                    if p.style is None or getattr(p.style, 'name', '') != 'Caption':
                        for r in p.runs:
                            r.italic = True
                except Exception:
                    pass
            continue

        # Blank line ends a paragraph buffer.
        if not line.strip():
            finalize_paragraph_buf(document, para_buf)
            i += 1
            continue

        # Fenced code block start
        mcode = re.match(r"^(\s*)(`{3,}|~{3,})(\w+)?\s*$", line)
        if mcode:
            finalize_paragraph_buf(document, para_buf)
            fence = mcode.group(2)
            language = mcode.group(3)
            i += 1
            code_lines: List[str] = []
            while i < len(lines) and not re.match(rf"^\s*{re.escape(fence)}\s*$", lines[i]):
                code_lines.append(lines[i])
                i += 1
            # Skip closing fence
            if i < len(lines):
                i += 1
            add_code_block(document, code_lines, language)
            continue

        # Table detection
        t_end, rows, aligns = parse_table_block(lines, i)
        if t_end > i:
            finalize_paragraph_buf(document, para_buf)
            if rows:
                nrows = len(rows)
                ncols = len(rows[0])
                table = document.add_table(rows=nrows, cols=ncols)
                # Apply a readable table style if available
                try:
                    table.style = "Light List Accent 1"
                except Exception:
                    pass
                # Fill cells with inline formatting (bold, inline code)
                for r_idx, row in enumerate(rows):
                    cells = table.rows[r_idx].cells
                    for c_idx, cell_text in enumerate(row):
                        cell = cells[c_idx]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        _add_inline_formatted_runs(p, cell_text)
            i = t_end
            continue

        # Heading
        mhead = _HEADING_RE.match(line)
        if mhead:
            finalize_paragraph_buf(document, para_buf)
            level = len(mhead.group(1))
            raw_title = mhead.group(2).strip()
            add_heading_with_bookmark(document, raw_title, level, next_bookmark_id)
            next_bookmark_id += 1
            current_heading_level = level
            i += 1
            continue

        # Image-only line
        mimg = _IMAGE_RE.fullmatch(line.strip())
        if mimg:
            finalize_paragraph_buf(document, para_buf)
            add_image(document, mimg.group("src"), base_dir)
            i += 1
            continue

        # Lists (bulleted/ordered)
        mb = _LIST_BULLET_RE.match(line)
        mo = _LIST_ORDERED_RE.match(line)
        if mb or mo:
            finalize_paragraph_buf(document, para_buf)
            if mb:
                indent = mb.group("indent")
                text = mb.group("text")
                level = list_level_for_indent(indent)
                p = document.add_paragraph(style="List Bullet")
                _add_inline_formatted_runs(p, text)
            else:
                indent = mo.group("indent")
                text = mo.group("text")
                level = list_level_for_indent(indent)
                p = document.add_paragraph(style="List Number")
                _add_inline_formatted_runs(p, text)
            # Indent to reflect nesting level
            try:
                p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
                # Use a hanging indent so wrapped lines align cleanly
                p.paragraph_format.first_line_indent = -Inches(0.25)
            except Exception:
                pass
            i += 1
            continue

        # Inline images within a paragraph: break the paragraph buffer and inject image.
        inline_imgs = list(_IMAGE_RE.finditer(line))
        if inline_imgs:
            # Flush preceding paragraph text (before first image) into buf
            cursor = 0
            for m in inline_imgs:
                pre_text = line[cursor : m.start()].strip()
                if pre_text:
                    para_buf.append(pre_text)
                    finalize_paragraph_buf(document, para_buf)
                add_image(document, m.group("src"), base_dir)
                cursor = m.end()
            post_text = line[cursor:].strip()
            if post_text:
                para_buf.append(post_text)
            i += 1
            continue

        # Default: accumulate paragraph text
        para_buf.append(line)
        i += 1

    finalize_paragraph_buf(document, para_buf)


def build_document(md_path: Path, out_path: Path, apply_template_style: bool = True) -> None:

    document = Document()

    # Title page is handled by the template if present. Insert a TOC after any
    # initial content (e.g., before main body when no specialized template).
    document.add_paragraph("Table of Contents", style="Heading 1")
    add_toc(document)

    md_text = md_path.read_text(encoding="utf-8")
    # Remove HTML comments outside fenced code blocks
    md_text = _strip_html_comments_outside_code(md_text)
    parse_markdown_to_docx(document, md_text, base_dir=md_path.parent)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    # Save the original, unmodified styling version
    document.save(str(out_path))

    # Additionally save a styled copy when a Word template is available
    if WORD_TEMPLATE.exists() and apply_template_style:
        styled_path = out_path.with_name(out_path.stem + "-styled.docx")
        shutil.copyfile(out_path, styled_path)
        to_copy = [
            "word/styles.xml",
        ]
        replace_docx_parts_from_template(WORD_TEMPLATE, styled_path, to_copy)


def export_docx(input_md: Path | None = None, out_path: Path | None = None) -> None:
    """Export spec.md (or input_md) to DOCX under exports/."""
    md_path = DEFAULT_MD if input_md is None else input_md
    output_path = OUTPUT_DOCX if out_path is None else out_path

    if not md_path.exists():
        raise FileNotFoundError(f"Input markdown not found: {md_path}")

    print(f"Generating: {output_path}")
    build_document(md_path, output_path)
    if WORD_TEMPLATE.exists():
        styled_path = output_path.with_name(output_path.stem + "-styled.docx")
        print(f"Generating: {styled_path}")
    print("Done.")


def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(description="Export DocLang spec markdown to DOCX")
    parser.add_argument(
        "--input",
        type=str,
        default=str(DEFAULT_MD),
        help="Path to input markdown (default: spec.md)",
    )

    args = parser.parse_args(argv)

    try:
        export_docx(Path(args.input))
    except FileNotFoundError as exc:
        print(str(exc), file=sys.stderr)
        return 2
    except Exception as e:
        print(f"Failed to write document: {e}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
