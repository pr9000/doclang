#!/usr/bin/env python3

"""
Generate an ISO draft DOCX from iso-standard.md.

Features:
- Determines output version X.Y.Z (defaults to latest X.Y and Z+1).
- Parses markdown headings, paragraphs, lists, code blocks, tables, images.
- Embeds images referenced in the markdown into the DOCX.
- Inserts a Word Table of Contents field (updates on open in Word).
- Optionally uses iso-standards/ISO_standard_template.dotx if present.

Dependencies:
- python-docx (pip install python-docx)

Note: This is a lightweight Markdown parser tailored for this repository's
iso-standard.md. It aims for faithful structure rather than exhaustive
Markdown feature coverage. If you need more fidelity, consider running
Pandoc externally and post-processing to add a TOC field.
"""

from __future__ import annotations

import argparse
import os
import shutil
import re
import sys
from pathlib import Path
from typing import List, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError as exc:
    print(
        "Missing dependency: python-docx. Install with `pip install python-docx`.",
        file=sys.stderr,
    )
    raise


ROOT = Path(__file__).resolve().parents[1]
ISO_STANDARDS_DIR = ROOT / "iso-standards"
ISO_TEMPLATE = ISO_STANDARDS_DIR / "ISO_standard_template.dotx"
DEFAULT_MD = ROOT / "iso-standard.md"


VERSION_FILENAME_RE = re.compile(r"^doctags_draft_v(\d+)\.(\d+)\.(\d+)\.docx$")


def find_latest_version(dir_path: Path) -> Optional[Tuple[int, int, int]]:
    latest: Optional[Tuple[int, int, int]] = None
    if not dir_path.exists():
        return None
    for name in os.listdir(dir_path):
        m = VERSION_FILENAME_RE.match(name)
        if not m:
            continue
        x, y, z = map(int, m.groups())
        if latest is None or (x, y, z) > latest:
            latest = (x, y, z)
    return latest


def compute_version(x: Optional[int], y: Optional[int], z: Optional[int]) -> Tuple[int, int, int]:
    latest = find_latest_version(ISO_STANDARDS_DIR)
    if x is None and y is None and z is None:
        # Use latest X.Y and increment Z. Default to 0.0.1 when no prior drafts.
        if latest is None:
            return (0, 0, 1)
        lx, ly, lz = latest
        return (lx, ly, lz + 1)

    # Some values provided. Fill missing ones sensibly.
    if x is None or y is None:
        if latest is None:
            # Start new series if no latest exists.
            x = x if x is not None else 0
            y = y if y is not None else 0
        else:
            lx, ly, _ = latest
            x = x if x is not None else lx
            y = y if y is not None else ly
    if z is None:
        # When X/Y are explicitly set, start at Z=0; otherwise continue from latest.
        if latest is None:
            z = 0
        else:
            lx, ly, lz = latest
            if x == lx and y == ly:
                z = lz + 1
            else:
                z = 0
    return (int(x), int(y), int(z))


def out_path_for_version(x: int, y: int, z: int) -> Path:
    return ISO_STANDARDS_DIR / f"doctags_draft_v{x}.{y}.{z}.docx"


def add_toc(document: Document) -> None:
    # Insert a Word ToC field: TOC \o "1-3" \h \z \u
    p = document.add_paragraph()
    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instr)

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_separate)

    # Placeholder text shown before updating fields in Word.
    hint = p.add_run("Table of Contents – update to populate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    hint._r.append(fld_end)


def finalize_paragraph_buf(document: Document, buf: List[str]) -> None:
    if not buf:
        return
    text = " ".join(line.strip() for line in buf).strip()
    if text:
        add_formatted_paragraph(document, text)
    buf.clear()


def replace_docx_parts_from_template(template_path: Path, docx_path: Path, parts: List[str]) -> None:
    """
    Copy specific parts from a DOCX/DOTX template into a target DOCX, replacing
    any existing entries with the same names. This avoids duplicate entries and
    the associated zipfile UserWarning messages.

    - template_path: path to a .docx/.dotx file to read parts from
    - docx_path: path to the .docx file to modify in-place
    - parts: list of ZIP member names to copy (e.g., 'word/styles.xml')
    """
    import zipfile

    parts_set = set(parts)
    tmp_path = docx_path.with_suffix(docx_path.suffix + ".tmp")

    # Build a new zip: copy everything from the destination except the parts
    # we intend to replace, then write the parts from the template.
    with zipfile.ZipFile(docx_path, "r") as zread, zipfile.ZipFile(tmp_path, "w") as zwrite:
        # Copy existing entries except those to be replaced
        for info in zread.infolist():
            if info.filename in parts_set:
                continue
            data = zread.read(info.filename)
            # Preserve metadata by reusing ZipInfo when writing
            zwrite.writestr(info, data)

        # Append parts from the template if present
        with zipfile.ZipFile(template_path, "r") as zsrc:
            for name in parts:
                try:
                    src_info = zsrc.getinfo(name)
                    data = zsrc.read(name)
                    zwrite.writestr(src_info, data)
                except KeyError:
                    # Part not present in template; skip
                    pass

    # Atomically replace the original docx with the rebuilt archive
    tmp_path.replace(docx_path)


_LIST_BULLET_RE = re.compile(r"^(?P<indent>\s*)([-+*])\s+(?P<text>.+?)\s*$")
_LIST_ORDERED_RE = re.compile(r"^(?P<indent>\s*)(\d+)([.)])\s+(?P<text>.+?)\s*$")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*?)\s*$")
_IMAGE_RE = re.compile(r'!\[(?P<alt>[^\]]*)\]\((?P<src>[^)\s]+)(?:\s+"(?P<title>[^"]*)")?\)')


def parse_table_block(lines: List[str], start: int) -> Tuple[int, List[List[str]], Optional[List[str]]]:
    # Detect a GitHub-style pipe table starting at index start.
    # Returns (end_index_exclusive, rows, aligns)
    table_lines = []
    i = start
    while i < len(lines) and "|" in lines[i].strip():
        table_lines.append(lines[i].strip())
        i += 1
        # Stop if we hit a blank line.
        if i < len(lines) and not lines[i].strip():
            break
    if len(table_lines) < 2:
        return start, [], None
    # Second line must be the separator with dashes and optional colons.
    header = [c.strip() for c in split_md_row(table_lines[0])]
    sep = [c.strip() for c in split_md_row(table_lines[1])]
    if not all(re.match(r"^:?-{3,}:?$", s) for s in sep) or len(sep) != len(header):
        return start, [], None

    aligns: List[str] = []
    for s in sep:
        if s.startswith(":") and s.endswith(":"):
            aligns.append("center")
        elif s.endswith(":"):
            aligns.append("right")
        elif s.startswith(":"):
            aligns.append("left")
        else:
            aligns.append("left")

    rows = [header]
    for r in table_lines[2:]:
        cells = [c.strip() for c in split_md_row(r)]
        # Normalize cell count
        while len(cells) < len(header):
            cells.append("")
        rows.append(cells[: len(header)])
    return i, rows, aligns


def split_md_row(row: str) -> List[str]:
    # Splits a pipe table row into cells, ignoring outer pipes.
    s = row.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.replace("\\|", "|") for c in re.split(r"\s*\|\s*", s)]


def add_code_block(document: Document, code_lines: List[str], language: Optional[str]) -> None:
    # Add a monospaced, preformatted code block.
    p = document.add_paragraph()
    # Light grey background for the whole block
    try:
        _set_paragraph_shading(p, fill="F2F2F2")
    except Exception:
        pass
    run = p.add_run()
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    for idx, line in enumerate(code_lines):
        if idx:
            run.add_break(WD_BREAK.LINE)
        run.add_text(line.rstrip("\n"))


INLINE_CODE_TAG_RE = re.compile(r"<inline-code>(.*?)</inline-code>")
BOLD_RE = re.compile(r"\*\*(.*?)\*\*")
CODE_SPAN_RE = re.compile(r"(<inline-code>(.*?)</inline-code>|`([^`]+)`)")


def _set_paragraph_shading(paragraph, fill: str = "F2F2F2") -> None:
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _set_run_shading(run, fill: str = "F2F2F2") -> None:
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    rPr.append(shd)


def _add_bold_and_text_runs(paragraph, text: str) -> None:
    # Split by bold markers and add runs accordingly
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            t = text[pos : m.start()]
            if t:
                paragraph.add_run(t)
        bold_txt = m.group(1)
        br = paragraph.add_run(bold_txt)
        br.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_inline_formatted_runs(paragraph, text: str) -> None:
    # First, split into code and non-code spans
    cursor = 0
    for m in CODE_SPAN_RE.finditer(text):
        if m.start() > cursor:
            normal_text = text[cursor : m.start()]
            if normal_text:
                _add_bold_and_text_runs(paragraph, normal_text)
        code_txt = m.group(2) if m.group(2) is not None else m.group(3)
        r = paragraph.add_run(code_txt)
        r.font.name = "Consolas"
        try:
            _set_run_shading(r, fill="F2F2F2")
        except Exception:
            pass
        cursor = m.end()
    if cursor < len(text):
        tail = text[cursor:]
        _add_bold_and_text_runs(paragraph, tail)


def add_formatted_paragraph(document: Document, text: str):
    p = document.add_paragraph()
    _add_inline_formatted_runs(p, text)
    return p


def add_image(document: Document, src: str, base_dir: Path) -> bool:
    # Resolve and add image; return True if added.
    # Skip if remote URL.
    if re.match(r"^[a-z]+://", src):
        return False
    img_path = (base_dir / src).resolve()
    if not img_path.exists():
        # Try relative to project root as fallback.
        alt_path = (ROOT / src).resolve()
        if alt_path.exists():
            img_path = alt_path
        else:
            print(f"[warn] Image not found: {src}", file=sys.stderr)
            return False
    try:
        # Resize to visible page width while preserving aspect ratio.
        try:
            section = document.sections[-1]
            max_width = section.page_width - section.left_margin - section.right_margin
            document.add_picture(str(img_path), width=max_width)
        except Exception:
            # Fallback if section metrics are unavailable
            document.add_picture(str(img_path))
        return True
    except Exception as e:
        print(f"[warn] Failed to embed image {src}: {e}", file=sys.stderr)
        return False


def parse_markdown_to_docx(document: Document, md_text: str, base_dir: Path) -> None:
    lines = md_text.splitlines()
    i = 0
    para_buf: List[str] = []

    # For list indentation, 2 spaces per level is common in this repo.
    def list_level_for_indent(s: str) -> int:
        return max(0, len(s.replace("\t", "  ")) // 2)

    while i < len(lines):
        line = lines[i]

        # Blank line ends a paragraph buffer.
        if not line.strip():
            finalize_paragraph_buf(document, para_buf)
            i += 1
            continue

        # Fenced code block start
        mcode = re.match(r"^(\s*)(`{3,}|~{3,})(\w+)?\s*$", line)
        if mcode:
            finalize_paragraph_buf(document, para_buf)
            fence = mcode.group(2)
            language = mcode.group(3)
            i += 1
            code_lines: List[str] = []
            while i < len(lines) and not re.match(rf"^\s*{re.escape(fence)}\s*$", lines[i]):
                code_lines.append(lines[i])
                i += 1
            # Skip closing fence
            if i < len(lines):
                i += 1
            add_code_block(document, code_lines, language)
            continue

        # Table detection
        t_end, rows, aligns = parse_table_block(lines, i)
        if t_end > i:
            finalize_paragraph_buf(document, para_buf)
            if rows:
                nrows = len(rows)
                ncols = len(rows[0])
                table = document.add_table(rows=nrows, cols=ncols)
                # Apply a readable table style if available
                try:
                    table.style = "Light List Accent 1"
                except Exception:
                    pass
                # Fill cells with inline formatting (bold, inline code)
                for r_idx, row in enumerate(rows):
                    cells = table.rows[r_idx].cells
                    for c_idx, cell_text in enumerate(row):
                        cell = cells[c_idx]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        _add_inline_formatted_runs(p, cell_text)
            i = t_end
            continue

        # Heading
        mhead = _HEADING_RE.match(line)
        if mhead:
            finalize_paragraph_buf(document, para_buf)
            level = len(mhead.group(1))
            text = mhead.group(2).strip()
            # python-docx supports heading levels 0..9; use 1..6 for Markdown.
            document.add_heading(text, level=level)
            i += 1
            continue

        # Image-only line
        mimg = _IMAGE_RE.fullmatch(line.strip())
        if mimg:
            finalize_paragraph_buf(document, para_buf)
            add_image(document, mimg.group("src"), base_dir)
            i += 1
            continue

        # Lists (bulleted/ordered)
        mb = _LIST_BULLET_RE.match(line)
        mo = _LIST_ORDERED_RE.match(line)
        if mb or mo:
            finalize_paragraph_buf(document, para_buf)
            if mb:
                indent = mb.group("indent")
                text = mb.group("text")
                level = list_level_for_indent(indent)
                p = document.add_paragraph(style="List Bullet")
                _add_inline_formatted_runs(p, text)
            else:
                indent = mo.group("indent")
                text = mo.group("text")
                level = list_level_for_indent(indent)
                p = document.add_paragraph(style="List Number")
                _add_inline_formatted_runs(p, text)
            # Indent to reflect nesting level
            try:
                p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
                # Use a hanging indent so wrapped lines align cleanly
                p.paragraph_format.first_line_indent = -Inches(0.25)
            except Exception:
                pass
            i += 1
            continue

        # Inline images within a paragraph: break the paragraph buffer and inject image.
        inline_imgs = list(_IMAGE_RE.finditer(line))
        if inline_imgs:
            # Flush preceding paragraph text (before first image) into buf
            cursor = 0
            for m in inline_imgs:
                pre_text = line[cursor : m.start()].strip()
                if pre_text:
                    para_buf.append(pre_text)
                    finalize_paragraph_buf(document, para_buf)
                add_image(document, m.group("src"), base_dir)
                cursor = m.end()
            post_text = line[cursor:].strip()
            if post_text:
                para_buf.append(post_text)
            i += 1
            continue

        # Default: accumulate paragraph text
        para_buf.append(line)
        i += 1

    finalize_paragraph_buf(document, para_buf)


def build_document(md_path: Path, out_path: Path, copy_iso_style: bool = True) -> None:

    document = Document()

    # Title page is handled by the template if present. Insert a TOC after any
    # initial content (e.g., before main body when no specialized template).
    document.add_paragraph("Table of Contents", style="Heading 1")
    add_toc(document)

    md_text = md_path.read_text(encoding="utf-8")
    parse_markdown_to_docx(document, md_text, base_dir=md_path.parent)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    # Save the original, unmodified styling version
    document.save(str(out_path))

    # Additionally save a copy with ISO styling applied, if requested and available
    if ISO_TEMPLATE.exists() and copy_iso_style:
        iso_styled_path = out_path.with_name(out_path.stem + ".iso-styling.docx")
        # Start from the original file contents
        shutil.copyfile(out_path, iso_styled_path)
        # Lift styles (and related assets) from the template into the copied docx,
        # replacing existing entries to avoid duplicate ZIP names warnings.
        to_copy = [
            "word/styles.xml",           # styles (required)
            # "word/numbering.xml",        # lists/bullets (if present)
            # "word/theme/theme1.xml",     # theme (if present)
            # "word/fontTable.xml",        # font table (optional)
        ]
        replace_docx_parts_from_template(ISO_TEMPLATE, iso_styled_path, to_copy)

def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(description="Write doctags ISO draft DOCX from markdown")
    parser.add_argument("--x", type=int, default=None, help="Version X (major)")
    parser.add_argument("--y", type=int, default=None, help="Version Y (minor)")
    parser.add_argument("--z", type=int, default=None, help="Version Z (patch)")
    parser.add_argument(
        "--input",
        type=str,
        default=str(DEFAULT_MD),
        help="Path to input markdown (default: iso-standard.md)",
    )

    args = parser.parse_args(argv)
    x, y, z = compute_version(args.x, args.y, args.z)
    out_path = out_path_for_version(x, y, z)

    print(out_path)
    
    md_path = Path(args.input)
    if not md_path.exists():
        print(f"Input markdown not found: {md_path}", file=sys.stderr)
        return 2

    print(f"Generating: {out_path}")
    try:
        build_document(md_path, out_path)
    except Exception as e:
        print(f"Failed to write document: {e}", file=sys.stderr)
        return 1
    # Inform about the ISO-styled file when template is available
    if ISO_TEMPLATE.exists():
        iso_styled_path = out_path.with_name(out_path.stem + ".iso-styling.docx")
        print(f"Also generated: {iso_styled_path}")
    print("Done.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
